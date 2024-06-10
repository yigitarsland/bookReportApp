from docx import Document
import requests
import re
import matplotlib.pyplot as plt
from PIL import Image
import io
from docx.shared import Inches


# URLs
book_url = "https://www.gutenberg.org/ebooks/68283.txt.utf-8"
picture_url = "https://static.wikia.nocookie.net/lovecraft/images/6/69/Cthulhu_Mythos_Hrairoo.png/revision/latest?cb=20221220005905"
logo_url = "https://i.etsystatic.com/21950283/r/il/3d34a3/2199302903/il_fullxfull.2199302903_1t20.jpg"

# Download the book
response = requests.get(book_url)
book_text = response.text

# Extract title, author's name, first chapter
title_search = re.search(r"Title:\s*(.*)", book_text)
title = title_search.group(1).strip() if title_search else "Title not found"
author_search = re.search(r"Author:\s*(.*)", book_text)
author = author_search.group(1).strip() if author_search else "Author not found"
first_chapter_search = re.search(r"(_1\. The Horror in Clay\._.*?)(_2\.|Chapter\s+2|End of the Project Gutenberg EBook)", book_text, re.DOTALL)
first_chapter_content = first_chapter_search.group(1).strip() if first_chapter_search else "First chapter not found"
print("Title:", title)
print("Author:", author)
print("\nFirst Chapter Content:\n")
print(first_chapter_content[:2000])  # Print the first 2000 characters of the first chapter for brevity

# Count the number of words in each paragraph
paragraphs = first_chapter_content.split('\n')
word_counts = [len(paragraph.split()) for paragraph in paragraphs]

# Print the word counts for each paragraph
for i, count in enumerate(word_counts):
    print(f"Paragraph {i+1}: {count} words")

# Create a distribution of paragraph lengths
paragraph_length_distribution = {}
for count in word_counts:
    if count in paragraph_length_distribution:
        paragraph_length_distribution[count] += 1
    else:
        paragraph_length_distribution[count] = 1

# Plot the distribution of paragraph lengths
plt.bar(paragraph_length_distribution.keys(), paragraph_length_distribution.values(), color='cyan')
plt.xlabel('Length of Paragraphs (Number of Words)')
plt.ylabel('Number of Paragraphs')
plt.title('Distribution of Paragraph Lengths in First Chapter')
plt.savefig('paragraph_lengths.png')  # Save the plot as an image file
plt.close()  # Close the plot to free up resources

# Download and process picture #1
response = requests.get(picture_url)
picture1 = Image.open(io.BytesIO(response.content))
cropped_picture1 = picture1.crop((20, 20, picture1.width - 20, picture1.height - 20))
resized_picture1 = cropped_picture1.resize((400, 400))
resized_picture1.save('picture1_processed.jpg')

# Download and process picture #2
response = requests.get(logo_url)
logo = Image.open(io.BytesIO(response.content))
logo = logo.convert("RGBA")  # Ensure the logo has an alpha channel
rotated_logo = logo.rotate(45, expand=True).resize((100,100))
rotated_logo.save('logo_rotated.png')

# Combine logo and picture #1
combined_image = resized_picture1.copy()
combined_image.paste(rotated_logo, (300, 300), rotated_logo)
combined_image.save('combined_image.png')

# Create a Word document
doc = Document()

# Title page
doc.add_heading(title, level=1)
doc.add_heading('by ' + author, level=2)
doc.add_heading('Project by Yigit Arslan', level=2)
doc.add_picture('combined_image.png', width=Inches(4.0))

doc.add_page_break()

# Info page
doc.add_heading('Distribution of Paragraph Lengths in First Chapter', level=1)
doc.add_picture('paragraph_lengths.png', width=Inches(5.5))

# Text paragraph including description of the plot
doc.add_paragraph(
    f'The first chapter of "{title}" by {author} contains {len(paragraphs)} paragraphs. '
    f'The distribution of paragraph lengths is shown in the plot above. '
    f'The paragraphs have a minimum of {min(word_counts)} words and a maximum of {max(word_counts)} words. '
    f'The average number of words per paragraph is {sum(word_counts) / len(word_counts):.2f}.'
)

doc.save('Lab11_Report.docx')
